from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "public" / "templates"
PDF_PATH = OUT_DIR / "reid-poole-music-booking-packet-template.pdf"
DOCX_PATH = OUT_DIR / "reid-poole-music-booking-packet-template.docx"

COLORS = {
    "night": "#0B0C09",
    "pine": "#263F38",
    "alpine": "#4F7870",
    "gold": "#C9962E",
    "sun": "#E0B861",
    "ivory": "#F4E7B8",
    "cream": "#F3DFA6",
    "copper": "#B86432",
    "ink": "#17201C",
    "mist": "#F9F2DF",
}

PLACEHOLDERS = [
    ("Client name", "{{CLIENT_NAME}}"),
    ("Event date", "{{EVENT_DATE}}"),
    ("Venue", "{{VENUE_NAME}}"),
    ("Venue address", "{{VENUE_ADDRESS}}"),
    ("Performance time", "{{PERFORMANCE_TIME}}"),
    ("Ensemble size", "{{ENSEMBLE_SIZE}}"),
    ("Total fee", "{{TOTAL_FEE}}"),
    ("Deposit amount", "{{DEPOSIT_AMOUNT}}"),
    ("Remaining balance", "{{REMAINING_BALANCE}}"),
    ("Special notes", "{{SPECIAL_NOTES}}"),
]

DETAIL_ROWS = [
    ("Client / planner", "{{CLIENT_NAME}} / {{PLANNER_NAME}}"),
    ("Email / phone", "{{CLIENT_EMAIL}} / {{CLIENT_PHONE}}"),
    ("Event date", "{{EVENT_DATE}}"),
    ("Venue", "{{VENUE_NAME}}"),
    ("Venue address", "{{VENUE_ADDRESS}}"),
    ("Performance window", "{{PERFORMANCE_TIME}}"),
    ("Ensemble", "{{ENSEMBLE_SIZE}}"),
    ("Sound provided by", "{{SOUND_PROVIDER}}"),
]

PAYMENT_ROWS = [
    ("Total performance fee", "{{TOTAL_FEE}}"),
    ("Deposit due to reserve date", "{{DEPOSIT_AMOUNT}}"),
    ("Remaining balance", "{{REMAINING_BALANCE}}"),
    ("Balance due date", "{{BALANCE_DUE_DATE}}"),
    ("Payment method", "{{PAYMENT_METHOD}}"),
]

TERMS = [
    "Date is reserved after agreement signature and deposit receipt.",
    "Client provides safe performance area, reasonable load-in access, and any required parking or vendor credentials.",
    "Outdoor performances require weather protection, shade when appropriate, and a stable performance surface.",
    "Custom arrangements and special song requests are quoted separately unless included in writing.",
    "Cancellation, postponement, and severe-weather changes should be handled in writing as early as possible.",
    "Final schedule, announcements, and music notes should be confirmed two weeks before the event when possible.",
]

QUESTION_SECTIONS = [
    ("Ceremony music", [
        "Processional song(s): ______________________________________________",
        "Bride / main entrance song: ________________________________________",
        "Recessional song: _________________________________________________",
        "Special moments or cultural traditions: ____________________________",
    ]),
    ("Cocktail hour / reception", [
        "Preferred feel: relaxed / upbeat / elegant / dance-forward",
        "Must-play songs: _________________________________________________",
        "Do-not-play notes: _______________________________________________",
        "Announcements or cues: ___________________________________________",
    ]),
    ("Venue logistics", [
        "Planner or day-of contact: ________________________________________",
        "Load-in time and entrance: ________________________________________",
        "Power location / stage size: ______________________________________",
        "Parking, green room, meals, or vendor notes: _______________________",
    ]),
]

SET_LISTS = [
    ("East Coast Swing / Cocktail Set", [
        "Ain't Misbehavin'",
        "All of Me",
        "Blue Bossa",
        "Fly Me to the Moon",
        "I Got Rhythm",
        "It Don't Mean a Thing",
        "On the Sunny Side of the Street",
    ]),
    ("Funk Brass Band / Party Set", [
        "Cissy Strut",
        "Chameleon",
        "Feel Like Funkin' It Up",
        "I Wish",
        "Pass the Peas",
        "September",
        "Street Parade",
    ]),
]

FAQS = [
    ("Can we request songs?", "Yes. Standard requests are welcome. Custom arrangements require advance notice and may include an arranging fee."),
    ("Can you play outdoors?", "Yes, when weather protection, shade, power, and a stable performance surface are provided."),
    ("Do you bring sound?", "Small sound can often be included. Larger rooms, outdoor stages, or full-band events may require production support."),
    ("When should final details be due?", "Two weeks before the event is ideal so the music plan, schedule, and logistics are clean."),
]

CHECKLIST = [
    "Signed agreement returned",
    "Deposit received",
    "Venue contact confirmed",
    "Schedule and performance windows confirmed",
    "Song requests finalized",
    "Load-in, parking, and power confirmed",
    "Remaining balance scheduled",
]


def hex_color(value: str) -> colors.HexColor:
    return colors.HexColor(value)


def pdf_on_page(canvas, doc):
    width, height = LETTER
    canvas.saveState()
    canvas.setFillColor(hex_color(COLORS["mist"]))
    canvas.rect(0, 0, width, height, fill=1, stroke=0)
    canvas.setStrokeColor(hex_color(COLORS["gold"]))
    canvas.setLineWidth(1.5)
    canvas.rect(0.42 * inch, 0.42 * inch, width - 0.84 * inch, height - 0.84 * inch, fill=0, stroke=1)
    canvas.setStrokeColor(hex_color(COLORS["alpine"]))
    canvas.setLineWidth(0.5)
    canvas.line(0.8 * inch, 0.68 * inch, width - 0.8 * inch, 0.68 * inch)
    canvas.setFillColor(hex_color(COLORS["pine"]))
    canvas.setFont("Helvetica-Bold", 8)
    canvas.drawString(0.8 * inch, 0.48 * inch, "REID POOLE MUSIC")
    canvas.setFillColor(hex_color(COLORS["alpine"]))
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(width - 0.8 * inch, 0.48 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CoverEyebrow",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=hex_color(COLORS["sun"]),
        spaceAfter=20,
    ))
    styles.add(ParagraphStyle(
        name="CoverTitle",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=38,
        leading=44,
        alignment=TA_CENTER,
        textColor=hex_color(COLORS["ivory"]),
        spaceAfter=14,
    ))
    styles.add(ParagraphStyle(
        name="CoverSub",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=14,
        leading=21,
        alignment=TA_CENTER,
        textColor=hex_color(COLORS["cream"]),
    ))
    styles.add(ParagraphStyle(
        name="H1",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=23,
        leading=29,
        textColor=hex_color(COLORS["pine"]),
        spaceBefore=6,
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name="H2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        textColor=hex_color(COLORS["copper"]),
        spaceBefore=12,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="BodyClean",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=14,
        textColor=hex_color(COLORS["ink"]),
        spaceAfter=7,
    ))
    styles.add(ParagraphStyle(
        name="Small",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=12,
        textColor=hex_color(COLORS["pine"]),
    ))
    styles.add(ParagraphStyle(
        name="SmallIvory",
        parent=styles["Small"],
        textColor=hex_color(COLORS["ivory"]),
        fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="Callout",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=15,
        textColor=hex_color(COLORS["ivory"]),
        alignment=TA_LEFT,
    ))

    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        leftMargin=0.78 * inch,
        rightMargin=0.78 * inch,
        topMargin=0.78 * inch,
        bottomMargin=0.78 * inch,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=pdf_on_page)])

    story = []
    story.append(Spacer(1, 0.95 * inch))
    story.append(Table(
        [
            [Paragraph("PREMIER BOOKING PACKET", styles["CoverEyebrow"])],
            [Paragraph("REID POOLE MUSIC", styles["CoverTitle"])],
            [Paragraph("Performance Agreement + Wedding Music Planning Questionnaire", styles["CoverSub"])],
            [Paragraph("Elegant, editable templates for weddings, private events, venues, festivals, and corporate bookings.", styles["CoverSub"])],
        ],
        colWidths=[doc.width],
        style=[
            ("BACKGROUND", (0, 0), (-1, -1), hex_color(COLORS["night"])),
            ("BOX", (0, 0), (-1, -1), 2, hex_color(COLORS["gold"])),
            ("INNERPADDING", (0, 0), (-1, -1), 28),
            ("TOPPADDING", (0, 0), (-1, -1), 18),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 18),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ],
    ))
    story.append(PageBreak())

    story.append(Paragraph("How to use this packet", styles["H1"]))
    story.append(Paragraph(
        "Duplicate this template for each client. Replace the bracketed placeholders, remove unused sections, then export the final copy as a PDF for signature or review.",
        styles["BodyClean"],
    ))
    placeholder_table = Table(
        [[Paragraph("<b>Field</b>", styles["Small"]), Paragraph("<b>Placeholder</b>", styles["Small"])]] +
        [[Paragraph(label, styles["Small"]), Paragraph(code, styles["Small"])] for label, code in PLACEHOLDERS],
        colWidths=[2.1 * inch, 4.1 * inch],
        style=[
            ("BACKGROUND", (0, 0), (-1, 0), hex_color(COLORS["pine"])),
            ("TEXTCOLOR", (0, 0), (-1, 0), hex_color(COLORS["ivory"])),
            ("GRID", (0, 0), (-1, -1), 0.35, hex_color(COLORS["alpine"])),
            ("BACKGROUND", (0, 1), (-1, -1), colors.white),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ],
    )
    story.append(placeholder_table)
    story.append(Spacer(1, 16))
    story.append(Table(
        [[Paragraph("Client-ready language", styles["Callout"])]],
        colWidths=[doc.width],
        style=[
            ("BACKGROUND", (0, 0), (-1, -1), hex_color(COLORS["pine"])),
            ("BOX", (0, 0), (-1, -1), 0.75, hex_color(COLORS["gold"])),
            ("LEFTPADDING", (0, 0), (-1, -1), 12),
            ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ("TOPPADDING", (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ],
    ))
    story.append(Paragraph(
        "Thank you for considering Reid Poole Music. This packet keeps the creative details, logistics, and agreement terms in one polished place so the event feels calm, personal, and well-planned.",
        styles["BodyClean"],
    ))
    story.append(PageBreak())

    story.append(Paragraph("Performance Agreement", styles["H1"]))
    story.append(Paragraph("Client and event details", styles["H2"]))
    story.append(_pdf_table(DETAIL_ROWS, doc.width, styles))
    story.append(Paragraph("Payment terms", styles["H2"]))
    story.append(_pdf_table(PAYMENT_ROWS, doc.width, styles))
    story.append(Paragraph("Core terms", styles["H2"]))
    story.append(_pdf_bullets(TERMS, styles))
    story.append(Spacer(1, 10))
    story.append(_pdf_signature_table(doc.width, styles))
    story.append(PageBreak())

    story.append(Paragraph("Wedding Music Planning Questionnaire", styles["H1"]))
    for title, prompts in QUESTION_SECTIONS:
        story.append(KeepTogether([
            Paragraph(title, styles["H2"]),
            _pdf_bullets(prompts, styles),
        ]))
    story.append(PageBreak())

    story.append(Paragraph("Sample Set Lists", styles["H1"]))
    set_tables = []
    for title, songs in SET_LISTS:
        set_tables.append([
            Paragraph(f"<b>{title}</b>", styles["SmallIvory"]),
            Paragraph("<br/>".join(songs), styles["Small"]),
        ])
    story.append(Table(set_tables, colWidths=[2.35 * inch, 3.85 * inch], style=[
        ("BACKGROUND", (0, 0), (0, -1), hex_color(COLORS["pine"])),
        ("TEXTCOLOR", (0, 0), (0, -1), hex_color(COLORS["ivory"])),
        ("BACKGROUND", (1, 0), (1, -1), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, hex_color(COLORS["alpine"])),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.append(Paragraph("Song request guide", styles["H2"]))
    story.append(_pdf_bullets([
        "Send must-play songs as early as possible.",
        "Include artist names or reference links to avoid confusion.",
        "Custom arrangements are most successful with four or more weeks of lead time.",
        "A short do-not-play list is helpful for weddings and corporate events.",
    ], styles))
    story.append(PageBreak())

    story.append(Paragraph("FAQ + Final Checklist", styles["H1"]))
    for question, answer in FAQS:
        story.append(Paragraph(question, styles["H2"]))
        story.append(Paragraph(answer, styles["BodyClean"]))
    story.append(Paragraph("Final checklist", styles["H2"]))
    story.append(_pdf_bullets(CHECKLIST, styles))

    doc.build(story)


def _pdf_table(rows, width, styles):
    data = [[Paragraph(label, styles["SmallIvory"]), Paragraph(value, styles["Small"])] for label, value in rows]
    return Table(data, colWidths=[2.05 * inch, width - 2.05 * inch], style=[
        ("BACKGROUND", (0, 0), (0, -1), hex_color(COLORS["pine"])),
        ("TEXTCOLOR", (0, 0), (0, -1), hex_color(COLORS["ivory"])),
        ("BACKGROUND", (1, 0), (1, -1), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, hex_color(COLORS["alpine"])),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ])


def _pdf_bullets(items, styles):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["BodyClean"]), bulletColor=hex_color(COLORS["gold"])) for item in items],
        bulletType="bullet",
        leftIndent=18,
    )


def _pdf_signature_table(width, styles):
    rows = [
        ("Client signature", "____________________________________  Date: __________"),
        ("Artist / representative", "____________________________________  Date: __________"),
    ]
    return _pdf_table(rows, width, styles)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=COLORS["ivory"])
        set_cell_text(cells[1], value)
        shade_cell(cells[0], COLORS["pine"])
    doc.add_paragraph()
    return table


def add_docx_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Georgia"
        run.font.color.rgb = RGBColor.from_string(COLORS["pine"].replace("#", ""))
    return p


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(7)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("REID POOLE MUSIC\n")
    run.bold = True
    run.font.name = "Georgia"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(COLORS["pine"].replace("#", ""))
    sub = cover.add_run("Premier Booking Packet\n")
    sub.font.name = "Arial"
    sub.font.size = Pt(16)
    sub.font.color.rgb = RGBColor.from_string(COLORS["copper"].replace("#", ""))
    body = cover.add_run("Performance Agreement + Wedding Music Planning Questionnaire")
    body.font.name = "Arial"
    body.font.size = Pt(11)
    body.font.color.rgb = RGBColor.from_string(COLORS["alpine"].replace("#", ""))

    doc.add_paragraph("Editable template for weddings, private events, venues, festivals, and corporate bookings. Replace placeholders, remove unused sections, and export the final version as PDF or DOCX.")
    add_heading(doc, "Quick edit fields", 1)
    add_docx_table(doc, PLACEHOLDERS)

    doc.add_page_break()
    add_heading(doc, "Performance Agreement", 1)
    doc.add_paragraph("This agreement confirms the musical services, payment terms, planning notes, and performance logistics for {{CLIENT_NAME}}.")
    add_heading(doc, "Client and event details", 2)
    add_docx_table(doc, DETAIL_ROWS)
    add_heading(doc, "Payment and deposit terms", 2)
    add_docx_table(doc, PAYMENT_ROWS)
    add_heading(doc, "Terms and expectations", 2)
    add_docx_bullets(doc, TERMS)
    add_heading(doc, "Signatures", 2)
    add_docx_table(doc, [
        ("Client signature", "____________________________________  Date: __________"),
        ("Artist / representative", "____________________________________  Date: __________"),
    ])

    doc.add_page_break()
    add_heading(doc, "Wedding Music Planning Questionnaire", 1)
    for title, prompts in QUESTION_SECTIONS:
        add_heading(doc, title, 2)
        for prompt in prompts:
            doc.add_paragraph(prompt)
        doc.add_paragraph()

    add_heading(doc, "Sample Set Lists", 1)
    for title, songs in SET_LISTS:
        add_heading(doc, title, 2)
        add_docx_bullets(doc, songs)

    add_heading(doc, "Song Request Guide", 1)
    add_docx_bullets(doc, [
        "Send must-play songs early with artist names or reference links.",
        "Custom arrangements may require additional lead time and arranging fees.",
        "Keep ceremony cues, announcements, and special moments in one final run-of-show.",
    ])

    doc.add_page_break()
    add_heading(doc, "FAQ", 1)
    for question, answer in FAQS:
        add_heading(doc, question, 2)
        doc.add_paragraph(answer)
    add_heading(doc, "Final Checklist", 1)
    add_docx_bullets(doc, CHECKLIST)
    add_heading(doc, "Special notes", 2)
    doc.add_paragraph("{{SPECIAL_NOTES}}")

    doc.core_properties.author = "Alpine Groove Guide"
    doc.core_properties.title = "Reid Poole Music Premier Booking Packet"
    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_pdf()
    build_docx()
    print(PDF_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
